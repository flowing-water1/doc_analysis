import json

import pandas as pd
import streamlit as st
import os
from utils import dataframe_agent, analysis_agent, word_similarity_agent
from docx import Document

os.environ["http_proxy"] = f"http://127.0.0.1:7890"
os.environ["https_proxy"] = f"http://127.0.0.1:7890"


def create_chart(input_data, chart_type):
    df_data = pd.DataFrame(input_data["data"], columns=input_data["columns"])
    df_data.set_index(input_data["columns"][0], inplace=True)
    if chart_type == "bar":
        st.bar_chart(df_data)
    elif chart_type == "line":
        st.line_chart(df_data)
    elif chart_type == "scatter":
        st.scatter_chart(df_data)


st.title("💡 数据分析智能工具")

with st.sidebar:
    openai_api_key = st.text_input("请输入OpenAI API密钥：", type="password")
    openai_api_base = st.text_input("请输入API_BASE：")
    st.markdown("[获取OpenAI API key](https://platform.openai.com/account/api-keys)")


data = st.file_uploader("上传你的数据文件（xlsx或csv格式）：", type=['xlsx', 'csv'])
if data:
    # 获取文件扩展名
    file_extension = data.name.split('.')[-1]

    # 根据扩展名选择读取方式
    if file_extension == 'xlsx':
        st.session_state["df"] = pd.read_excel(data)
    elif file_extension == 'csv':
        st.session_state["df"] = pd.read_csv(data)
    else:
        st.error("不支持的文件格式，请上传xlsx或csv文件。")

    with st.expander("原始数据"):
        st.dataframe(st.session_state["df"])

import streamlit as st
from docx import Document



query = st.text_area("请输入你关于以上表格的问题，或数据提取请求，或可视化要求（支持散点图、折线图、条形图）：")
button = st.button("生成回答", key="1")

st.divider()

st.title("📈表格数据分析")
if "df" not in st.session_state:
    st.info("直接在上面那里传入即可")
else:
    with st.expander("原始数据"):
        st.dataframe(st.session_state["df"])
analysis_button = st.button("分析数据", key="2")

st.divider()

st.title("📄 Word文件相似度比对")
# 创建两列
col1, col2 = st.columns(2)

# 第一列：处理第一个Word文件
with col1:
    word_data_1 = st.file_uploader("上传你的数据文件1（docx格式）：", type="docx")
    if word_data_1:
        # 读取Word内容
        doc = Document(word_data_1)
        st.session_state["word_content_1"] = "\n".join([para.text for para in doc.paragraphs])

        # 在页面显示Word内容
        with st.expander("原始数据1"):
            st.text(st.session_state["word_content_1"])

# 第二列：处理第二个Word文件
with col2:
    word_data_2 = st.file_uploader("上传你的数据文件2（docx格式）：", type="docx")
    if word_data_2:
        # 读取Word内容
        doc = Document(word_data_2)
        st.session_state["word_content_2"] = "\n".join([para.text for para in doc.paragraphs])

        # 在页面显示Word内容
        with st.expander("原始数据2"):
            st.text(st.session_state["word_content_2"])
word_button = st.button("比对相似度", key="3")

if (button or analysis_button or word_button) and not openai_api_key:
    st.toast("请先输入OpenAI API密钥", icon="🚨")

if (button or analysis_button) and "df" not in st.session_state:

    st.toast("请先上传excel数据文件",icon="🚨")

if word_button and "word_content_1" not in st.session_state and "word_content_2" not in st.session_state:
    st.toast("请先上传word数据文件",icon="🚨")


if button and openai_api_key and "df" in st.session_state:
    with st.spinner("AI正在思考中，请稍等..."):
        response, read = dataframe_agent(openai_api_key, openai_api_base, st.session_state["df"], query)
        st.write(response)
        st.divider()
        st.write("思考过程：")
        st.json(read, expanded=False)
        st.divider()

        # 检查 read 中的每个元素
        if isinstance(read, list):
            for i, element in enumerate(read):

                if isinstance(element, tuple) and len(element) > 0:
                    log_content = element[0].log if hasattr(element[0], 'log') else None
                    remaining_content = element[1] if len(element) > 1 else None

                    if log_content:
                        st.write(f"思考过程{i + 1}：")
                        st.write(log_content)

                    # 确保 remaining_content 是有效的类型
                    if remaining_content is not None:
                        # 检查是否为 Pandas Index 类型
                        if isinstance(remaining_content, pd.Index):
                            st.write(remaining_content.tolist())  # 转换为列表形式展示
                            st.divider()
                        else:
                            st.write(remaining_content)
                            st.divider()
                    else:
                        st.write(f"元素 {i} 中没有后续内容。")
        else:
            st.write("read 不是有效的列表格式。")

        response_dict = json.loads(response)
        if "answer" in response_dict:
            st.write(response_dict["answer"])
        if "table" in response_dict:
            st.table(pd.DataFrame(response_dict["table"]["data"],
                                  columns=response_dict["table"]["columns"]))
        if "bar" in response_dict:
            create_chart(response_dict["bar"], "bar")
        if "line" in response_dict:
            create_chart(response_dict["line"], "line")
        if "scatter" in response_dict:
            create_chart(response_dict["scatter"], "scatter")

if analysis_button and openai_api_key and "df" in st.session_state:
    with st.spinner("AI正在分析数据中，请稍等..."):
        analysis_result, read = analysis_agent(openai_api_key, openai_api_base, st.session_state["df"])
        st.write(analysis_result)
        st.divider()
        st.write("思考过程：")
        st.json(read, expanded=False)
        st.divider()
        # 检查 read 中的每个元素
        if isinstance(read, list):
            for i, element in enumerate(read):

                if isinstance(element, tuple) and len(element) > 0:
                    log_content = element[0].log if hasattr(element[0], 'log') else None
                    remaining_content = element[1] if len(element) > 1 else None

                    if log_content:
                        st.write(f"思考过程{i + 1}：")
                        st.write(log_content)

                    # 确保 remaining_content 是有效的类型
                    if remaining_content is not None:
                        # 检查是否为 Pandas Index 类型
                        if isinstance(remaining_content, pd.Index):
                            st.write(remaining_content.tolist())  # 转换为列表形式展示
                            st.divider()
                        else:
                            st.write(remaining_content)
                            st.divider()
                    else:
                        st.write(f"元素 {i} 中没有后续内容。")
        else:
            st.write("read 不是有效的列表格式。")

if word_button and openai_api_key and word_data_1 and word_data_2:
    with st.spinner("AI正在分析数据中，请稍等..."):
        word_result= word_similarity_agent(openai_api_key, openai_api_base,
                                                       st.session_state["word_content_1"],
                                                       st.session_state["word_content_2"])

        st.write(word_result)